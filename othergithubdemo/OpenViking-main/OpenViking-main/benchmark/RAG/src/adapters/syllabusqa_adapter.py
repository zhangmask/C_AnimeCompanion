# src/adapters/syllabusqa_adapter.py
"""
SyllabusQA Dataset Adapter

SyllabusQA is a syllabus QA dataset containing 39 syllabi and 5078 questions.
Each question is about a specific syllabus, with answer types including:
- single factual: single factual question
- multi factual: multi factual question
- single reasoning: single reasoning question
- multi reasoning: multi reasoning question
- summarization: summarization question
- yes/no: yes/no question
- no answer: no answer can be found in the syllabus

Dataset characteristics:
1. Each question has only one answer
2. Each answer has corresponding answer_span (evidence span)
3. reasoning type questions have reasoning steps
4. Original documents are in docx format

Adapter functions:
- data_prepare: Convert docx to Markdown format
- load_and_transform: Parse QA data, category stores question type
- build_prompt: Build QA prompt
- post_process_answer: Post-process LLM output
"""

import json
import os
import csv
from typing import List, Dict, Any

from .base import BaseAdapter, StandardDoc, StandardSample, StandardQA

# Rule for when answer cannot be found
MISSING_RULE = "If no information is available to answer the question, write 'Not mentioned'."

# Specific instructions for different categories
CATEGORY_INSTRUCTIONS = {
    "single factual": """Extract the single factual answer from the syllabus.
- Use EXACT wording from context when possible
- Provide concise, direct answer
- Do NOT add extra info or explanation""",
    
    "multi factual": """Extract multiple factual answers from the syllabus.
- Use EXACT wording from context when possible
- List items separated by commas
- Include all relevant facts""",
    
    "single reasoning": """Answer using simple logical reasoning based on the syllabus.
- Use ONLY facts from context
- Make clear, direct conclusion
- Do NOT explain reasoning
- Do NOT invent information""",
    
    "multi reasoning": """Answer using reasoning based on the syllabus.
- Use ONLY facts from context
- Do NOT invent information""",
    
    "summarization": """Summarize relevant information from the syllabus.
- Provide concise summary covering key points
- Use wording from syllabus when possible
- Include all important details""",
    
    "yes/no": """Answer Yes/No question based on the syllabus.
- First respond "Yes" or "No"
- Do NOT add explanation
- Use ONLY info from context
- Do NOT invent information"""
}


class SyllabusQAAdapter(BaseAdapter):
    """
    Adapter specifically for processing SyllabusQA dataset.
    
    Converts syllabi (docx) to Markdown documents,
    and converts QA data to standardized StandardSample format.
    
    Attributes:
        raw_file_path: Raw CSV data file path
        syllabus_dir: docx file directory path
        logger: Logger
    """
    
    def __init__(self, raw_file_path: str, **kwargs):
        """
        Initialize SyllabusQAAdapter.
        
        Args:
            raw_file_path: Raw data file path (CSV or merged JSON)
            **kwargs: Other parameters (ignored, for compatibility)
        """
        super().__init__(raw_file_path)
        # docx file directory, defaults to syllabi subdirectory under data directory
        if os.path.isdir(raw_file_path):
            base_dir = raw_file_path
        else:
            base_dir = os.path.dirname(raw_file_path)
        
        # Check for official repo structure first
        official_syllabus_dir = os.path.join(base_dir, 'syllabi', 'syllabi_redacted', 'word')
        if os.path.exists(official_syllabus_dir):
            self.syllabus_dir = official_syllabus_dir
        else:
            # Fallback to original structure
            self.syllabus_dir = os.path.join(base_dir, 'syllabi')
    
    def data_prepare(self, doc_dir: str) -> List[StandardDoc]:
        """
        Load raw docx files and convert to OpenViking-friendly format.
        
        Only process syllabus documents mentioned in CSV, avoid processing irrelevant documents.
        Requires python-docx library to parse docx files.
        
        Args:
            doc_dir: Document output directory path
            
        Returns:
            List[StandardDoc]: List of standardized document objects
            
        Raises:
            FileNotFoundError: syllabus directory not found
        """
        if not os.path.exists(self.syllabus_dir):
            raise FileNotFoundError(f"Syllabus directory not found: {self.syllabus_dir}")

        res: List[StandardDoc] = []
        os.makedirs(doc_dir, exist_ok=True)
        
        # Get list of syllabus_name mentioned in CSV
        required_syllabi = self._get_required_syllabi()
        self.logger.info(f"[SyllabusQAAdapter] Required syllabi from CSV: {len(required_syllabi)}")
        
        # Get all docx files
        docx_files = [f for f in os.listdir(self.syllabus_dir) if f.endswith('.docx')]
        
        for docx_file in docx_files:
            syllabus_id = docx_file.replace('.docx', '')
            
            # Only process syllabi mentioned in CSV
            if syllabus_id not in required_syllabi:
                continue
            
            docx_path = os.path.join(self.syllabus_dir, docx_file)
            
            try:
                # Convert docx to Markdown
                doc_content = self._convert_docx_to_markdown(docx_path)
                
                doc_path = os.path.join(doc_dir, f"{syllabus_id}_doc.md")
                with open(doc_path, "w", encoding="utf-8") as f:
                    f.write(doc_content)
                res.append(StandardDoc(syllabus_id, doc_path))
            except Exception as e:
                self.logger.error(f"[syllabusqa adapter] doc:{syllabus_id} prepare error {e}")
                # If python-docx not installed, try using plain text
                if "No module named 'docx'" in str(e):
                    self.logger.warning("python-docx not installed, skipping docx conversion")
                    break
                raise e
        
        self.logger.info(f"[SyllabusQAAdapter] Processed {len(res)} syllabus documents")
        return res

    def _get_required_syllabi(self) -> set:
        """
        Get list of syllabus_name mentioned in CSV or JSON.
        
        Returns:
            set: syllabus_name set
        """
        required = set()
        
        # Determine data source type
        if self.raw_file_path.endswith('.json'):
            # Load from JSON
            if not os.path.exists(self.raw_file_path):
                return required
            
            with open(self.raw_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            for syllabus_name in data.keys():
                if syllabus_name:
                    required.add(syllabus_name)
        elif self.raw_file_path.endswith('.csv'):
            csv_files = [self.raw_file_path]
        elif os.path.isdir(self.raw_file_path):
            csv_files = [os.path.join(self.raw_file_path, f) 
                        for f in os.listdir(self.raw_file_path) 
                        if f.endswith('.csv')]
        else:
            return required
        
        # Process CSV files if any
        if 'csv_files' in locals():
            for csv_file in csv_files:
                if not os.path.exists(csv_file):
                    continue
                with open(csv_file, 'r', encoding='utf-8') as f:
                    import csv
                    reader = csv.DictReader(f)
                    for row in reader:
                        syllabus_name = row.get('syllabus_name', '')
                        if syllabus_name:
                            required.add(syllabus_name)
        
        return required

    def _convert_docx_to_markdown(self, docx_path: str) -> str:
        """
        Convert docx file to Markdown string.
        
        Args:
            docx_path: docx file path
            
        Returns:
            str: Markdown formatted content
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        doc = Document(docx_path)
        md_lines = []
        
        # Extract filename as title
        filename = os.path.basename(docx_path).replace('.docx', '')
        md_lines.append(f"# {filename}")
        md_lines.append("")
        
        # Iterate through all paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Check if heading style
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level_num = int(level)
                        md_lines.append(f"{'#' * level_num} {text}")
                    except ValueError:
                        md_lines.append(f"## {text}")
                else:
                    md_lines.append(text)
                md_lines.append("")
        
        # Extract tables
        for table in doc.tables:
            md_lines.append("## Table")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
            md_lines.append("")
        
        return "\n".join(md_lines)

    def load_and_transform(self) -> List[StandardSample]:
        """
        Load raw CSV data and convert to standardized StandardSample object list.
        
        Processing logic:
        1. Read CSV file (supports single CSV or directory)
        2. Group by syllabus_name
        3. Format question as "Based on the syllabus \"{syllabus_name}\", {question}"
        4. category stores question_type
        5. answer_span as evidence
        
        Returns:
            List[StandardSample]: List of standardized sample objects
            
        Raises:
            FileNotFoundError: Raw data file not found
        """
        # Determine if CSV file or JSON file
        if self.raw_file_path.endswith('.json'):
            return self._load_from_json()
        elif self.raw_file_path.endswith('.csv'):
            return self._load_from_csv([self.raw_file_path])
        elif os.path.isdir(self.raw_file_path):
            # Directory, find all CSV files
            csv_files = [os.path.join(self.raw_file_path, f) 
                        for f in os.listdir(self.raw_file_path) 
                        if f.endswith('.csv')]
            return self._load_from_csv(csv_files)
        else:
            raise FileNotFoundError(f"Unsupported file format: {self.raw_file_path}")

    def _load_from_json(self) -> List[StandardSample]:
        """
        Load data from JSON file.
        """
        if not os.path.exists(self.raw_file_path):
            raise FileNotFoundError(f"Raw data file not found: {self.raw_file_path}")

        with open(self.raw_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        standard_samples = []

        for syllabus_name, qa_list in data.items():
            qa_pairs = []
            
            for qa_item in qa_list:
                question = qa_item.get("question", "")
                answer = qa_item.get("answer", "")
                question_type = qa_item.get("question_type", "")
                qa_id = qa_item.get("id", "")
                
                # Skip "no answer" type questions as RAG results cannot be evaluated
                if question_type == "no answer":
                    continue
                
                # Collect answer_span as evidence
                evidence = []
                for i in range(1, 6):
                    span = qa_item.get(f"answer_span_{i}", "")
                    if span and span.strip():
                        evidence.append(span.strip())
                
                # Collect reasoning_steps, also as evidence (for reasoning type questions)
                reasoning_steps = []
                for i in range(1, 6):
                    step = qa_item.get(f"reasoning_step_{i}", "")
                    if step and step.strip():
                        reasoning_steps.append(step.strip())
                        # reasoning_steps also added to evidence for recall calculation
                        if step.strip() not in evidence:
                            evidence.append(step.strip())
                
                # Format question
                formatted_question = f'Based on the syllabus "{syllabus_name}", {question}'
                
                qa_pairs.append(StandardQA(
                    question=formatted_question,
                    gold_answers=[answer] if answer else ["Not mentioned"],
                    evidence=evidence,
                    category=question_type,
                    metadata={
                        "id": qa_id,
                        "reasoning_steps": reasoning_steps
                    }
                ))

            # Only add samples with QA pairs to result
            if qa_pairs:
                standard_samples.append(StandardSample(
                    sample_id=syllabus_name,
                    qa_pairs=qa_pairs
                ))

        return standard_samples

    def _load_from_csv(self, csv_files: List[str]) -> List[StandardSample]:
        """
        Load data from CSV files.
        
        Args:
            csv_files: List of CSV file paths
            
        Returns:
            List[StandardSample]: List of standardized sample objects
        """
        # Group by syllabus_name
        syllabus_qa_map: Dict[str, List] = {}
        
        for csv_file in csv_files:
            if not os.path.exists(csv_file):
                self.logger.warning(f"CSV file not found: {csv_file}")
                continue
            
            with open(csv_file, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    syllabus_name = row.get('syllabus_name', '')
                    if syllabus_name not in syllabus_qa_map:
                        syllabus_qa_map[syllabus_name] = []
                    syllabus_qa_map[syllabus_name].append(row)
        
        standard_samples = []
        
        for syllabus_name, qa_list in syllabus_qa_map.items():
            qa_pairs = []
            
            for qa_item in qa_list:
                question = qa_item.get("question", "")
                answer = qa_item.get("answer", "")
                question_type = qa_item.get("question_type", "")
                qa_id = qa_item.get("id", "")
                
                # Skip "no answer" type questions as RAG results cannot be evaluated
                if question_type == "no answer":
                    continue
                
                # Collect answer_span as evidence
                evidence = []
                for i in range(1, 6):
                    span = qa_item.get(f"answer_span_{i}", "")
                    if span and span.strip():
                        evidence.append(span.strip())
                
                # Collect reasoning_steps, also as evidence (for reasoning type questions)
                reasoning_steps = []
                for i in range(1, 6):
                    step = qa_item.get(f"reasoning_step_{i}", "")
                    if step and step.strip():
                        reasoning_steps.append(step.strip())
                        # reasoning_steps also added to evidence for recall calculation
                        if step.strip() not in evidence:
                            evidence.append(step.strip())
                
                # Format question
                formatted_question = f'Based on the syllabus "{syllabus_name}", {question}'
                
                qa_pairs.append(StandardQA(
                    question=formatted_question,
                    gold_answers=[answer] if answer else ["Not mentioned"],
                    evidence=evidence,
                    category=question_type,
                    metadata={
                        "id": qa_id,
                        "reasoning_steps": reasoning_steps
                    }
                ))

            # Only add samples with QA pairs to result
            if qa_pairs:
                standard_samples.append(StandardSample(
                    sample_id=syllabus_name,
                    qa_pairs=qa_pairs
                ))

        return standard_samples

    def build_prompt(self, qa: StandardQA, context_blocks: List[str]) -> tuple[str, Dict[str, Any]]:
        """
        Build complete prompt to send to LLM.
        
        Prompt structure:
        1. Context content (retrieved document fragments)
        2. Category-specific instructions
        3. Rule for when answer cannot be found
        4. Question
        
        Args:
            qa: Standardized QA object
            context_blocks: List of retrieved context text blocks
            
        Returns:
            tuple[str, Dict[str, Any]]: 
                - Complete prompt string
                - Metadata dictionary, containing id
        """
        eff_q = qa.question
        category = qa.category
        
        category_instruction = CATEGORY_INSTRUCTIONS.get(category, "")
        
        context_text = "\n\n".join(context_blocks)
        
        if category_instruction:
            full_prompt = f"{context_text}\n\n{category_instruction}\n\n{MISSING_RULE}\n\nQuestion: {eff_q}\n\nAnswer:"
        else:
            full_prompt = f"{context_text}\n\n{MISSING_RULE}\n\nQuestion: {eff_q}\n\nAnswer:"

        meta = {"id": qa.metadata.get("id", "")}
        return full_prompt, meta

    def post_process_answer(self, qa: StandardQA, raw_answer: str, meta: Dict[str, Any]) -> str:
        """
        Post-process raw answer generated by LLM.
        
        Current implementation only strips leading/trailing whitespace.
        
        Args:
            qa: Standardized QA object
            raw_answer: Raw answer generated by LLM
            meta: Metadata dictionary
            
        Returns:
            str: Processed answer
        """
        return raw_answer.strip()
